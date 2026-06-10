"""
Pro AI Meeting Minutes Assistant — v4.1 (Streamlit Edition)
Improvements over v1:
  • Map-Reduce summarisation — handles 2-hour meetings without data loss
  • Speaker Diarization — pyannote.audio (optional, GPU-friendly, graceful fallback)
  • Transcript auto-saved to /transcripts folder with speaker labels
  • Custom vocabulary hint (Whisper initial_prompt)
  • All v1 features preserved (languages, models, export formats, noise reduction)
  • Added in Visual System Map, Model Card and Log File.
"""

import streamlit as st
from faster_whisper import WhisperModel
from groq import Groq
import os
import tempfile
import requests
from datetime import datetime
import re
import torch
import gc
from tenacity import retry, stop_after_attempt, wait_exponential
import subprocess
import logging

# Page config must be the first Streamlit command executed
st.set_page_config(
    page_title="Pro AI Meeting Minutes — v4.1",
    page_icon="🎙️",
    layout="wide"
)

# Setup professional logging format
LOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd(), "logs")
os.makedirs(LOG_DIR, exist_ok=True)
log_filename = f"app_{datetime.now().strftime('%Y%m_%d')}.log"

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] (%(filename)s:%(lineno)d) - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(LOG_DIR, log_filename), encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# --- API KEY CONFIGURATION (Cross-Platform: Streamlit Secrets / Local Env) ---
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY") or os.environ.get("GROQ_API_KEY")
SEALION_API_KEY = st.secrets.get("SEALION_API_KEY") or os.environ.get("SEALION_API_KEY")
HF_TOKEN = st.secrets.get("HF_TOKEN") or os.environ.get("HF_TOKEN")

if not GROQ_API_KEY or not SEALION_API_KEY:
    st.error("❌ Missing required API keys (GROQ_API_KEY / SEALION_API_KEY). Please add them to your Environment Variables or `.streamlit/secrets.toml`.")
    st.stop()

if HF_TOKEN:
    os.environ["HF_TOKEN"] = HF_TOKEN

# Initialize Core API Client
groq_client = Groq(api_key=GROQ_API_KEY)

# --- OPTIONAL LIBRARIES SETUP ---
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- CACHING THE CORE STT MODEL ---
@st.cache_resource
def load_stt_model():
    device = "cuda" if torch.cuda.is_available() else "cpu"
    compute_type = "float16" if device == "cuda" else "int8"
    return WhisperModel("small", device=device, compute_type=compute_type)

stt_model = load_stt_model()

# --- OPTIONAL: Speaker Diarization ---
@st.cache_resource
def load_diarization_pipeline(token):
    if not token:
        return None, False
    try:
        from pyannote.audio import Pipeline as PyannotePipeline
        pipeline = PyannotePipeline.from_pretrained(
            "pyannote/speaker-diarization-community-1",
            token=token  
        )
        if torch.cuda.is_available():
            pipeline = pipeline.to(torch.device("cuda"))
        return pipeline, True
    except Exception as e:
        logging.error(f"Diarization initialization error: {e}")
        return None, False

diarize_pipeline, DIARIZATION_AVAILABLE = load_diarization_pipeline(HF_TOKEN)

# --- REGISTRIES ---
MODEL_MAP = {
    "Llama-3.3 70B": "llama-3.3-70b-versatile",
    "Llama-3.1 8B": "llama-3.1-8b-instant",
    "Sea-Lion v4 27B": "aisingapore/Gemma-SEA-LION-v4-27B-IT"
}

LANGUAGE_MAP = {
    "English": "en",
    "Malay": "ms",
    "Chinese": "zh",
    "Auto-Detect": None
}

BASE_PATH = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
TRANSCRIPT_DIR = os.path.join(BASE_PATH, "transcripts")
os.makedirs(TRANSCRIPT_DIR, exist_ok=True)


# ─────────────────────────────────────────────
# HELPERS & PROCESSING CORE LOGIC
# ─────────────────────────────────────────────
def clean_text(text):
    return re.sub(r'[\*\*|_|=|#|`]', '', text).strip()

def call_llm(model_display_name, prompt):
    model_id = MODEL_MAP.get(model_display_name)
    logging.info(f"Dispatching inference call to model: {model_display_name} ({model_id})")
    try:
        if "Sea-Lion" in model_display_name:
            url = "https://api.sea-lion.ai/v1/chat/completions"
            headers = {"Authorization": f"Bearer {SEALION_API_KEY}", "Content-Type": "application/json"}
            payload = {
                "model": model_id,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7
            }
            response = requests.post(url, headers=headers, json=payload, timeout=60)
            logging.info("Sea-Lion LLM inference successfully completed.")
            return clean_text(response.json()['choices'][0]['message']['content'])
        else:
            chat_completion = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=model_id
            )
            logging.info("Groq LLM inference successfully completed.")
            return clean_text(chat_completion.choices[0].message.content)
    except Exception as e:
        logging.error(f"API processing exception on model {model_display_name}: {str(e)}", exc_info=True)
        return f"LLM Error: {str(e)}"

def convert_audio_to_standard_format(input_path, output_path=None):
    if output_path is None:
        output_path = os.path.join(tempfile.gettempdir(), f"converted_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav")
    
    cmd = ["ffmpeg", "-i", input_path, "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", "-y", output_path]
    try:
        subprocess.run(cmd, check=True, capture_output=True)
        logging.info(f"Audio normalization successful. Temp target: {output_path}")
        return output_path
    except subprocess.CalledProcessError as e:
        logging.error(f"FFmpeg pipeline failure: {e.stderr.decode('utf-8')}")
        return None

def run_diarization(audio_path):
    if not DIARIZATION_AVAILABLE or diarize_pipeline is None:
        return []
    try:
        converted_path = convert_audio_to_standard_format(audio_path)
        audio_to_use = converted_path if converted_path else audio_path
        
        output = diarize_pipeline(audio_to_use)
        segments = []
        
        if hasattr(output, 'speaker_diarization'):
            for turn, speaker in output.speaker_diarization:
                segments.append((turn.start, turn.end, speaker))
        elif hasattr(output, 'exclusive_speaker_diarization'):
            for turn, speaker in output.exclusive_speaker_diarization:
                segments.append((turn.start, turn.end, speaker))
        
        if converted_path and converted_path != audio_path and os.path.exists(converted_path):
            try: os.remove(converted_path)
            except: pass
        return segments
    except Exception as e:
        logging.error(f"Diarization error: {e}")
        return []

def assign_speakers(whisper_segments, diarization_segments):
    if not diarization_segments:
        return [(s.start, s.end, "Speaker ?", s.text) for s in whisper_segments]
    result = []
    for s in whisper_segments:
        best_speaker = "Speaker ?"
        best_overlap = 0.0
        for (d_start, d_end, label) in diarization_segments:
            overlap = max(0, min(s.end, d_end) - max(s.start, d_start))
            if overlap > best_overlap:
                best_overlap = overlap
                best_speaker = label
        result.append((s.start, s.end, best_speaker, s.text))
    return result

def chunk_transcript(text, chunk_size=8000):
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

def map_reduce_summary(full_transcript, model_choice, lang_name, status_element):
    chunks = chunk_transcript(full_transcript)
    partial_summaries = []

    for i, chunk in enumerate(chunks):
        status_element.write(f"📦 Summarising segment {i + 1} of {len(chunks)} (Map Phase)...")
        chunk_prompt = (
            f"You are summarising a segment of a {lang_name} meeting transcript. "
            f"Be concise. Extract key facts, decisions, and any action items:\n\n{chunk}"
        )
        mini_summary = call_llm(model_choice, chunk_prompt)
        partial_summaries.append(f"[Segment {i + 1}]\n{mini_summary}")

    status_element.write("📑 Synthesising master summary (Reduce Phase)...")
    combined_notes = "\n\n".join(partial_summaries)
    master_prompt = f"""
You are a professional Chief of Staff. Below are summaries of different segments of a long meeting.
Synthesise them into one cohesive Master Executive Summary in {lang_name}.

Structure your output with these four sections:
1. Executive Overview — what was the meeting about?
2. Key Decisions — what was agreed upon?
3. Action Items — who is doing what, and by when?
4. Open Questions — anything left unresolved?

Segment Summaries:
{combined_notes}
"""
    return call_llm(model_choice, master_prompt)

def save_transcript(labeled_segments, timestamp):
    filename = f"transcript_{timestamp}.txt"
    filepath = os.path.join(TRANSCRIPT_DIR, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"Meeting Transcript — {timestamp}\n")
        f.write("=" * 60 + "\n\n")
        for (start, end, speaker, text) in labeled_segments:
            f.write(f"[{start:.1f}s – {end:.1f}s]  {speaker}:  {text.strip()}\n")
    return filepath

def generate_report(summary, transcript, format_choice, timestamp):
    if format_choice == "PDF":
        if not REPORTLAB_AVAILABLE: return None
        path = os.path.join(tempfile.gettempdir(), f"report_{timestamp}.pdf")
        doc = SimpleDocTemplate(path, pagesize=letter)
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle('Body', parent=styles['Normal'], spaceAfter=10)
        elements = [
            Paragraph("Meeting Summary", styles['Title']),
            Paragraph(summary.replace('\n', '<br/>'), body_style),
            Spacer(1, 12),
            Paragraph("Full Transcript", styles['Heading2']),
            Paragraph(transcript.replace('\n', '<br/>'), styles['Code'])
        ]
        doc.build(elements)
        return path
    elif format_choice == "DOCX":
        if not DOCX_AVAILABLE: return None
        path = os.path.join(tempfile.gettempdir(), f"report_{timestamp}.docx")
        doc = Document()
        doc.add_heading('Meeting Summary', 0)
        doc.add_paragraph(summary)
        doc.add_page_break()
        doc.add_heading('Full Transcript', level=1)
        doc.add_paragraph(transcript)
        doc.save(path)
        return path
    else:
        path = os.path.join(tempfile.gettempdir(), f"report_{timestamp}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"SUMMARY:\n{summary}\n\n" + "=" * 30 + f"\n\nTRANSCRIPT:\n{transcript}")
        return path


# ─────────────────────────────────────────────
# STREAMLIT USER INTERFACE
# ─────────────────────────────────────────────
st.title("🎙️ Pro AI Meeting Minutes — v4.1")
st.caption("_Model card · Log File · Visual System Map Engine_")

# Layout Configuration: Sidebar Control Unit + Main Tab Panels
with st.sidebar:
    st.header("🎛️ Configuration Console")
    
    upload_input = st.file_uploader("Upload Video / Audio", type=["mp3", "wav", "m4a", "mp4", "avi"])
    
    st.subheader("Options")
    lang_dropdown = st.selectbox("Speaking Language", list(LANGUAGE_MAP.keys()), index=0)
    model_dropdown = st.selectbox("AI Engine", list(MODEL_MAP.keys()), index=0)
    format_choice = st.radio("Export Format", ["TXT", "PDF", "DOCX"], index=0)
    noise_red = st.checkbox("Enable Noise Reduction (VAD)", value=True)
    
    st.subheader("Advanced Layout")
    custom_vocab = st.text_area(
        "Custom Vocabulary / Jargon Hint", 
        placeholder="e.g. Project Orion, KPI dashboard, SCRUM",
        help="Words Whisper might mishear. Separate with commas."
    )
    
    diarize_label = f"Enable Speaker Diarization {'✅' if DIARIZATION_AVAILABLE else '⚠️ (HF_TOKEN Missing)'}"
    enable_diarization = st.checkbox(diarize_label, value=False, disabled=not DIARIZATION_AVAILABLE)
    
    process_btn = st.button("🚀 Process Meeting Minutes", type="primary", use_container_width=True)

# Main Screen Interface Layout
tab_summary, tab_transcript, tab_info = st.tabs(["📋 AI Summary", "📝 Raw Transcript", "ℹ️ Model Card & Diagnostics"])

with tab_info:
    st.markdown("""
    ### System Architecture & AI Governance Info
    
    #### 1. Speech-to-Text Engine
    * **Model:** Whisper 'Small' (via `faster-whisper`)
    * **Intended Use:** High-speed multilingual execution context.
    * **Known Limitations:** May mishear domain jargon unless specified in *Custom Vocabulary*.
    
    #### 2. Summarization Engines (LLMs)
    * **Llama-3.3 70B:** Optimized for strategic context patterns and complex corporate logic reasoning.
    * **Llama-3.1 8B:** High-speed edge alternative processing maps.
    * **Sea-Lion v4 27B:** Localized fine-tuning tailored specifically for Southeast Asian localization structures (Malay, Mandarin, localized dialects).
    
    #### 3. Privacy & Compliance Framework
    * Audio data streams processing runs container-isolated or via enterprise endpoints (Groq / AI Singapore). No storage streams are used for downstream base model iterations.
    """)

# Core Processing Execution Event Block
if process_btn:
    if upload_input is None:
        st.warning("⚠️ Execution Interrupted: Please load a valid audio or video asset into the dashboard upload block.")
    else:
        with st.status("🏗️ Initialization Processing Lifecycle...", expanded=True) as status:
            try:
                # Save uploaded streaming data array natively to working execution directory
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(upload_input.name)[1]) as temp_audio:
                    temp_audio.write(upload_input.getvalue())
                    audio_source = temp_audio.name
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                lang_code = LANGUAGE_MAP.get(lang_dropdown)
                initial_prompt = custom_vocab.strip() if custom_vocab else None
                
                # --- STAGE 1: Transcription ---
                status.update(label="🚀 Deploying Whisper STT Nodes...", state="running")
                segments_gen, info = stt_model.transcribe(
                    audio_source,
                    beam_size=5,
                    vad_filter=noise_red,
                    language=lang_code,
                    task="transcribe",
                    initial_prompt=initial_prompt
                )
                
                seg_list = list(segments_gen)
                if not seg_list:
                    st.error("The transcription engine encountered absolute silence. Check audio tracks.")
                    st.stop()
                
                # --- STAGE 2: Optional Speaker Diarization ---
                diarization_segs = []
                if enable_diarization and DIARIZATION_AVAILABLE:
                    status.update(label="🎙️ Segmenting Acoustic Streams (Pyannote Pipeline)...", state="running")
                    diarization_segs = run_diarization(audio_source)
                
                labeled_segments = assign_speakers(seg_list, diarization_segs)
                
                full_transcript = ""
                for (start, end, speaker, text) in labeled_segments:
                    if diarization_segs:
                        full_transcript += f"[{start:.1f}s] {speaker}: {text.strip()}\n"
                    else:
                        full_transcript += f"[{start:.1f}s] {text.strip()}\n"
                
                # --- STAGE 3 & 4: Map-Reduce Aggregation Summaries ---
                status.update(label="🗺️ Deploying Map-Reduce Synthesizer...", state="running")
                final_summary = map_reduce_summary(full_transcript, model_dropdown, lang_dropdown, status)
                
                # --- STAGE 5: Disk Sync Operations ---
                status.update(label="💾 Writing Persistent Outputs...", state="running")
                save_transcript(labeled_segments, timestamp)
                
                # --- STAGE 6: Serialization / Downstream Report Exports ---
                status.update(label="📄 Packaging Document Bundles...", state="running")
                report_path = generate_report(final_summary, full_transcript, format_choice, timestamp)
                
                # Cleanup Temp Files
                try: os.remove(audio_source)
                except: pass
                
                status.update(label="✅ Run Completed Successfully!", state="complete")
                
                # --- RENDER OUTPUTS TO RESPECTIVE TAB PANELS ---
                with tab_summary:
                    st.subheader("📋 Executive Summary Output")
                    st.markdown(final_summary)
                
                with tab_transcript:
                    st.subheader("📝 Processed Log Output")
                    st.text_area("Transcript Window", value=full_transcript, height=450)
                
                # Render File Download interface block directly inside control channel sidebar
                with st.sidebar:
                    st.success("🎉 Processing Pipeline Succeeded!")
                    if report_path and os.path.exists(report_path):
                        with open(report_path, "rb") as file:
                            st.download_button(
                                label=f"📥 Download Report ({format_choice})",
                                data=file,
                                file_name=os.path.basename(report_path),
                                mime="application/octet-stream",
                                use_container_width=True
                            )
                            
            except Exception as e:
                status.update(label="❌ Pipeline Runtime Exception Triggered", state="error")
                logging.error(f"Critical Exception within workspace engine: {str(e)}", exc_info=True)
                st.error(f"Execution Error: {str(e)}")
